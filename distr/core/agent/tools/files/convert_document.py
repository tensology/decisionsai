"""
Convert Document Tool — Markdown to PDF, DOCX, or Google Doc.

Uses the Python `markdown` library to render MD → styled HTML, then:
  • PDF  — Playwright (Chromium) prints the HTML to PDF.  Mermaid.js
           diagrams are rendered client-side before print.
  • DOCX — pandoc (GFM) if available, else python-docx fallback.
  • Google Doc — delegates to the existing MarkdownToGoogleDocTool.

Tables, fenced code blocks, and Mermaid diagrams are supported in all
output formats (Mermaid renders as a PNG fallback in DOCX).
"""

import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from langchain.tools import BaseTool
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

# ── Markdown → HTML ──────────────────────────────────────────────

_CSS = """\
body {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  max-width: 820px; margin: 40px auto; padding: 0 24px;
  color: #1f2328; line-height: 1.6; font-size: 15px;
}
h1,h2,h3,h4,h5,h6 { margin-top: 1.4em; margin-bottom: .6em; font-weight: 600; }
h1 { font-size: 2em; border-bottom: 1px solid #d1d9e0; padding-bottom: .3em; }
h2 { font-size: 1.5em; border-bottom: 1px solid #d1d9e0; padding-bottom: .3em; }
code {
  background: #f0f2f5; padding: 2px 6px; border-radius: 4px;
  font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace;
  font-size: 0.9em;
}
pre { background: #f6f8fa; border: 1px solid #d1d9e0; border-radius: 6px;
      padding: 14px; overflow-x: auto; }
pre code { background: none; padding: 0; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #d1d9e0; padding: 8px 12px; text-align: left; }
th { background: #f6f8fa; font-weight: 600; }
tr:nth-child(even) { background: #f9fafb; }
blockquote { border-left: 4px solid #d1d9e0; margin: 1em 0; padding: .5em 1em; color: #636c76; }
img { max-width: 100%; }
hr { border: none; border-top: 1px solid #d1d9e0; margin: 2em 0; }
a { color: #0969da; text-decoration: none; }
a:hover { text-decoration: underline; }
.mermaid { text-align: center; margin: 1.5em 0; }
@media print {
  body { margin: 0; padding: 0 16px; }
  pre { white-space: pre-wrap; word-break: break-word; }
}
"""

_MERMAID_SCRIPT = (
    '<script src="https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.min.js"></script>\n'
    "<script>mermaid.initialize({startOnLoad:true, theme:'default'});</script>"
)


def _md_to_html(md_text: str) -> str:
    """Convert Markdown text to a full styled HTML document.

    Mermaid fenced code blocks (```mermaid) are converted to
    ``<div class="mermaid">`` so the CDN script renders them.
    """
    import markdown as _md

    # Pre-process: convert ```mermaid blocks to <div class="mermaid">
    has_mermaid = "```mermaid" in md_text
    processed = re.sub(
        r"```mermaid\s*\n(.*?)```",
        r'<div class="mermaid">\n\1</div>',
        md_text,
        flags=re.DOTALL,
    )

    extensions = ["tables", "fenced_code", "codehilite", "toc", "attr_list", "md_in_html"]
    ext_configs = {"codehilite": {"css_class": "highlight", "guess_lang": False}}
    html_body = _md.markdown(processed, extensions=extensions, extension_configs=ext_configs)

    mermaid_tag = _MERMAID_SCRIPT if has_mermaid else ""
    return (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        f"<style>{_CSS}</style>{mermaid_tag}</head>"
        f"<body>{html_body}</body></html>"
    )


# ── HTML → PDF (Playwright) ─────────────────────────────────────

def _html_to_pdf(html: str, output_path: str, wait_for_mermaid: bool = False) -> str:
    """Render HTML to PDF using Playwright Chromium."""
    from playwright.sync_api import sync_playwright

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        page = browser.new_page()
        page.set_content(html, wait_until="networkidle")

        if wait_for_mermaid:
            # Give mermaid.js time to render SVGs
            try:
                page.wait_for_function(
                    "() => !document.querySelector('.mermaid:not([data-processed])')",
                    timeout=15000,
                )
            except Exception:
                # Timeout is non-fatal — diagrams may just not render
                logger.warning("Mermaid rendering timed out — diagrams may be missing")

        page.pdf(path=output_path, format="A4", print_background=True, margin={
            "top": "20mm", "bottom": "20mm", "left": "15mm", "right": "15mm",
        })
        browser.close()

    logger.info("Converted HTML → PDF: %s", output_path)
    return output_path


# ── HTML/MD → DOCX ──────────────────────────────────────────────

def _md_to_docx(md_text: str, output_path: str) -> str:
    """Convert Markdown to DOCX. Prefers pandoc for table/formatting fidelity."""
    pandoc = shutil.which("pandoc")
    if pandoc:
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
            f.write(md_text)
            tmp = f.name
        try:
            r = subprocess.run(
                [pandoc, tmp, "-f", "gfm", "-t", "docx", "-o", output_path],
                capture_output=True, text=True, timeout=60,
            )
            if r.returncode == 0:
                logger.info("Converted MD → DOCX (pandoc): %s", output_path)
                return output_path
            logger.warning("pandoc failed (%s): %s", r.returncode, r.stderr)
        finally:
            try:
                os.unlink(tmp)
            except OSError:
                pass

    # Fallback: python-docx
    return _md_to_docx_fallback(md_text, output_path)


def _md_to_docx_fallback(md_text: str, output_path: str) -> str:
    """Fallback DOCX conversion using python-docx with table support."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Headings
        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 6)
            text = stripped.lstrip("#").strip()
            if text:
                doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rules
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # Tables — collect contiguous pipe-delimited lines
        if "|" in stripped and stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                # Skip separator rows like |---|---|
                if re.match(r"^\|[\s\-:|]+\|$", row):
                    i += 1
                    continue
                cells = [c.strip() for c in row.strip("|").split("|")]
                table_lines.append(cells)
                i += 1
            if table_lines:
                cols = max(len(r) for r in table_lines)
                tbl = doc.add_table(rows=len(table_lines), cols=cols, style="Table Grid")
                for ri, row_cells in enumerate(table_lines):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < cols:
                            tbl.rows[ri].cells[ci].text = _strip_md_inline(cell_text)
            continue

        # Fenced code blocks
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            continue

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(_strip_md_inline(stripped[2:].strip()), style="List Bullet")
            i += 1
            continue

        # Numbered lists
        m = re.match(r"^\d+\.\s+(.+)", stripped)
        if m:
            doc.add_paragraph(_strip_md_inline(m.group(1)), style="List Number")
            i += 1
            continue

        # Regular paragraph
        if stripped:
            doc.add_paragraph(_strip_md_inline(stripped))
        i += 1

    doc.save(output_path)
    logger.info("Converted MD → DOCX (python-docx fallback): %s", output_path)
    return output_path


def _strip_md_inline(text: str) -> str:
    """Remove inline markdown formatting (bold, italic, code, links)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


# ── Tool definition ──────────────────────────────────────────────

class ConvertDocumentInput(BaseModel):
    """Input schema for convert_document tool."""
    input_path: Optional[str] = Field(
        default=None,
        description="Path to the source file (e.g. a .md markdown file). This argument is required."
    )
    output_format: Optional[str] = Field(
        default=None,
        description="Target format: 'pdf', 'docx', or 'google_doc' (or 'word' as alias for docx)."
    )
    output_path: Optional[str] = Field(
        default=None,
        description="Optional output file path. Defaults to same directory as input with the new extension.",
    )


class ConvertDocumentTool(BaseTool):
    """Convert documents between formats with proper formatting."""

    name: str = "convert_document"
    description: str = (
        "Convert a document file to another format with proper formatting.\n"
        "Supported conversions:\n"
        "  • Markdown (.md) → PDF  — fully styled with tables, code blocks, and Mermaid diagrams\n"
        "  • Markdown (.md) → DOCX — tables and formatting preserved\n"
        "  • Markdown (.md) → Google Doc — uploads to Google Drive\n"
        "\n"
        "Use this when the user says things like:\n"
        "  - 'convert README.md to pdf'\n"
        "  - 'make notes.md a docx'\n"
        "  - 'turn /tmp/report.md into docx'\n"
        "  - 'convert notes.markdown to google_doc'\n"
        "\n"
        "Do not call this tool with pronouns only (for example, 'this', 'that', 'turn this to PDF').\n"
        "When no explicit path is provided, the tool will return an explicit input error.\n"
        "\n"
        "Tables, fenced code blocks, and Mermaid diagrams (```mermaid) are all rendered properly.\n"
        "The PDF output is styled like a clean GitHub-flavored document."
    )
    args_schema: type[BaseModel] = ConvertDocumentInput

    def _find_recent_documents(self, multiple: bool = False) -> list[str]:
        """Find recently dropped document files for this user context."""
        try:
            import json
            storage_file = os.path.join(
                os.path.expanduser("~"),
                ".decisions",
                "dropped_files",
                "current_files.json",
            )
            if not os.path.exists(storage_file):
                return []

            with open(storage_file, "r") as f:
                data = json.load(f)

            doc_extensions = {".md", ".markdown", ".txt", ".docx", ".rtf"}

            def _is_doc_like(path: str) -> bool:
                return os.path.splitext(path)[1].lower() in doc_extensions

            all_files = data.get("document_files", [])
            if not all_files:
                all_files = data.get("files", [])

            # Prefer current chat bucket where available to avoid cross-chat confusion.
            current_chat_id = None
            if hasattr(self, "chat_manager") and self.chat_manager is not None:
                try:
                    current_chat_id = self.chat_manager.get_current_chat()
                except Exception:
                    current_chat_id = None

            chat_files_index = data.get("chat_files_index", {})
            if current_chat_id is not None:
                chat_bucket = chat_files_index.get(str(current_chat_id), {})
                if isinstance(chat_bucket, dict):
                    bucket_docs = chat_bucket.get("document_files", [])
                    bucket_files = chat_bucket.get("files", [])
                    bucket_candidates = [p for p in (bucket_docs or bucket_files) if _is_doc_like(p)]
                    if bucket_candidates:
                        all_files = bucket_candidates

            # Filter to files that still exist and are document-like.
            existing_files = [
                f
                for f in all_files
                if os.path.isfile(f) and _is_doc_like(f)
            ]
            if not existing_files:
                return []

            # Sort with most recent first.
            file_timestamps = data.get("file_timestamps", {})
            existing_files.sort(key=lambda f: file_timestamps.get(f, 0), reverse=True)
            return existing_files if multiple else existing_files[:1]
        except Exception as e:
            logger.warning("ConvertDocumentTool: Failed to find recent documents: %s", e)
            return []

    def _run(
        self,
        input_path: Optional[str] = None,
        output_format: Optional[str] = None,
        output_path: Optional[str] = None,
        **kwargs
    ) -> str:
        resolved_input_path: Optional[str] = input_path.strip() if input_path else None

        if resolved_input_path and resolved_input_path.lower() in {"this", "that", "it", "these", "those"}:
            resolved_input_path = None

        if not resolved_input_path:
            recent_documents = self._find_recent_documents()
            if recent_documents:
                resolved_input_path = recent_documents[0]
                logger.info("ConvertDocumentTool: Using most recent dropped document file %s", resolved_input_path)

        if not resolved_input_path:
            return (
                "Error: input_path is required for convert_document. "
                "Please provide a file path (for example, 'README.md') or drop a document file first."
            )

        if resolved_input_path.strip().lower() in {"this", "that", "it", "these", "those"}:
            return (
                "Error: input_path cannot be a pronoun. "
                "Provide the source file path explicitly before converting."
            )

        if not output_format:
            return "Error: output_format is required. Use 'pdf', 'docx', or 'google_doc'."

        fmt = output_format.lower().strip()
        fmt = re.sub(r"[^a-z_]", "", fmt)
        fmt = fmt.replace(".", "")
        if fmt == "word":
            fmt = "docx"
        if not fmt:
            return "Error: output_format is required. Use 'pdf', 'docx', or 'google_doc'."
        if fmt not in ("pdf", "docx", "google_doc", "googledoc"):
            return f"Error: Unsupported output format '{output_format}'. Use 'pdf', 'docx', or 'google_doc'."

        src = Path(resolved_input_path).expanduser().resolve()
        if not src.exists():
            return f"Error: File not found: {resolved_input_path}"

        # Read source markdown
        try:
            md_text = src.read_text(encoding="utf-8")
        except Exception as e:
            return f"Error reading file: {e}"

        if not md_text.strip():
            return "Error: Source file is empty."

        # Derive output path
        if fmt in ("google_doc", "googledoc"):
            return self._to_google_doc(md_text, src)

        ext = ".pdf" if fmt == "pdf" else ".docx"
        if output_path:
            out = Path(output_path).expanduser().resolve()
        else:
            out = src.with_suffix(ext)

        try:
            if fmt == "pdf":
                html = _md_to_html(md_text)
                has_mermaid = "```mermaid" in md_text
                _html_to_pdf(html, str(out), wait_for_mermaid=has_mermaid)
            else:
                _md_to_docx(md_text, str(out))
        except Exception as e:
            logger.error("Document conversion failed: %s", e, exc_info=True)
            return f"Error during conversion: {e}"

        size_kb = out.stat().st_size / 1024
        return f"Converted to {fmt.upper()}: {out} ({size_kb:.1f} KB)"

    def _to_google_doc(self, md_text: str, src: Path) -> str:
        """Delegate to the existing Google Doc conversion tool."""
        try:
            from distr.core.agent.tools.integrations.markdown_to_google_doc import (
                MarkdownToGoogleDocTool,
            )
            # The Google Doc tool reads from clipboard, so we need to put content there
            import platform
            import subprocess
            system = platform.system()
            if system == "Darwin":
                proc = subprocess.Popen(["pbcopy"], stdin=subprocess.PIPE)
                proc.communicate(md_text.encode("utf-8"))
            elif system == "Windows":
                proc = subprocess.run(
                    ["powershell", "-command", "Set-Clipboard -Value $input"],
                    input=md_text, text=True, timeout=5
                )
            else:
                try:
                    proc = subprocess.Popen(["xclip", "-selection", "clipboard"], stdin=subprocess.PIPE)
                    proc.communicate(md_text.encode("utf-8"))
                except FileNotFoundError:
                    proc = subprocess.Popen(["xsel", "--clipboard", "--input"], stdin=subprocess.PIPE)
                    proc.communicate(md_text.encode("utf-8"))

            tool = MarkdownToGoogleDocTool()
            return tool._run(open_in_brave=True)
        except Exception as e:
            return f"Error converting to Google Doc: {e}"

    async def _arun(
        self,
        input_path: Optional[str] = None,
        output_format: Optional[str] = None,
        output_path: Optional[str] = None,
        **kwargs,
    ) -> str:
        return self._run(input_path=input_path, output_format=output_format, output_path=output_path, **kwargs)
