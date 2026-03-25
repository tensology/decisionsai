"""
Document Extraction Tool for LangChain.

This tool extracts text content from various document formats:
- Plain text files (.txt, .md, .py, .js, .json, .xml, .html, .css, .csv, .log, .yaml, .yml, .ini, .cfg, .sh, .bat, etc.)
- PDF documents (.pdf)
- Word documents (.doc, .docx)
- Excel documents (.xls, .xlsx)
- ZIP archives (.zip)
- RAR archives (.rar, .rar5)
"""

from typing import Optional, List
from langchain.tools import BaseTool
from pydantic import BaseModel, Field
import logging
import os
import zipfile
import platform
import subprocess

logger = logging.getLogger(__name__)


class DocumentExtractorInput(BaseModel):
    """Input schema for document_extractor tool."""
    file_path: str = Field(description="Path to the document file to extract content from")
    extract_archives: bool = Field(default=False, description="If True, extract and process files from ZIP/RAR archives")


class DocumentExtractorTool(BaseTool):
    """Tool for extracting text content from various document formats."""
    
    name: str = "document_extractor"
    description: str = (
        "Extract and read text content from files. Use this tool to read ANY file. "
        "Supports: plain text files (.txt, .md, .py, .js, .json, .xml, .html, .css, .csv, .log, .yaml, .sh, etc.), "
        "PDF documents (.pdf), Word documents (.doc, .docx), "
        "Excel spreadsheets (.xls, .xlsx), ZIP archives (.zip), and RAR archives (.rar). "
        "Use this tool when the user wants to read a file, see file contents, or extract text from any document. "
        "This is the PRIMARY tool for reading files and documents."
    )
    args_schema: type[BaseModel] = DocumentExtractorInput
    
    # Text file extensions that can be read directly
    TEXT_EXTENSIONS: set = {
        '.txt', '.md', '.markdown', '.rst', '.text',
        '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.c', '.cpp', '.h', '.hpp', '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.scala',
        '.html', '.htm', '.xml', '.svg', '.css', '.scss', '.sass', '.less',
        '.json', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf', '.properties',
        '.sh', '.bash', '.zsh', '.fish', '.bat', '.cmd', '.ps1',
        '.sql', '.graphql', '.gql',
        '.csv', '.tsv', '.log',
        '.gitignore', '.dockerignore', '.env', '.env.example',
        '.makefile', '.dockerfile',
        ''  # Files without extension (like Makefile, Dockerfile)
    }
    
    def _extract_text_file(self, file_path: str, max_chars: int = 100000) -> str:
        """Extract text from a plain text file."""
        try:
            # Try UTF-8 first
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            except UnicodeDecodeError:
                # Fallback to latin-1 which can read any byte
                with open(file_path, 'r', encoding='latin-1') as f:
                    content = f.read()
            
            # Truncate if too long
            if len(content) > max_chars:
                content = content[:max_chars] + f"\n\n... [Content truncated at {max_chars} characters. File is {len(content)} characters total.]"
            
            return content
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}", exc_info=True)
            return f"Error reading text file: {str(e)}"
    
    def _is_text_file(self, file_path: str) -> bool:
        """Check if file is likely a text file based on extension or content."""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Check known text extensions
        if file_ext in self.TEXT_EXTENSIONS:
            return True
        
        # Check for extensionless files that are commonly text (like Makefile, Dockerfile)
        basename = os.path.basename(file_path).lower()
        text_basenames = {'makefile', 'dockerfile', 'gemfile', 'rakefile', 'procfile', 'readme', 'license', 'changelog', 'authors', 'contributing'}
        if basename in text_basenames:
            return True
        
        # Try to detect if it's a text file by reading first bytes
        try:
            with open(file_path, 'rb') as f:
                chunk = f.read(8192)
                # Check for null bytes which indicate binary
                if b'\x00' in chunk:
                    return False
                # Try to decode as UTF-8
                try:
                    chunk.decode('utf-8')
                    return True
                except UnicodeDecodeError:
                    # Try latin-1 as fallback
                    try:
                        chunk.decode('latin-1')
                        # If mostly printable ASCII, it's probably text
                        printable_ratio = sum(1 for b in chunk if 32 <= b <= 126 or b in (9, 10, 13)) / len(chunk)
                        return printable_ratio > 0.7
                    except Exception:
                        return False
        except Exception:
            return False
    
    def _extract_pdf(self, file_path_or_bytes) -> str:
        """Extract text from PDF file (supports file path or BytesIO)."""
        try:
            import io
            is_bytes = isinstance(file_path_or_bytes, (io.BytesIO, bytes))
            
            # Try pdfplumber first (better for complex PDFs)
            try:
                import pdfplumber
                text_parts = []
                if is_bytes:
                    if isinstance(file_path_or_bytes, bytes):
                        file_path_or_bytes = io.BytesIO(file_path_or_bytes)
                    with pdfplumber.open(file_path_or_bytes) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                else:
                    with pdfplumber.open(file_path_or_bytes) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text_parts.append(page_text)
                return "\n\n".join(text_parts)
            except ImportError:
                # Fallback to PyPDF2
                import PyPDF2
                text_parts = []
                if is_bytes:
                    if isinstance(file_path_or_bytes, bytes):
                        file_path_or_bytes = io.BytesIO(file_path_or_bytes)
                    pdf_reader = PyPDF2.PdfReader(file_path_or_bytes)
                    for page in pdf_reader.pages:
                        text_parts.append(page.extract_text())
                else:
                    with open(file_path_or_bytes, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text_parts.append(page.extract_text())
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}", exc_info=True)
            return f"Error extracting PDF: {str(e)}"
    
    def _extract_word(self, file_path_or_bytes) -> str:
        """Extract text from Word document (.doc or .docx). Supports file path or BytesIO."""
        try:
            import io
            is_bytes = isinstance(file_path_or_bytes, (io.BytesIO, bytes))
            
            # Check if it's .docx format
            is_docx = False
            if is_bytes:
                # For bytes, assume .docx (most common in archives)
                is_docx = True
            else:
                is_docx = file_path_or_bytes.lower().endswith('.docx')
            
            if is_docx:
                # Modern Word format - requires python-docx
                try:
                    from docx import Document
                except ImportError as e:
                    raise RuntimeError(
                        "python-docx is not installed. Install with: pip install python-docx"
                    ) from e
                if is_bytes:
                    if isinstance(file_path_or_bytes, bytes):
                        file_path_or_bytes = io.BytesIO(file_path_or_bytes)
                    doc = Document(file_path_or_bytes)
                else:
                    doc = Document(file_path_or_bytes)
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                return "\n".join(paragraphs)
            else:
                # Legacy .doc format - requires antiword or textutil (macOS)
                system = platform.system()
                if system == "Darwin":  # macOS
                    if is_bytes:
                        return "Error: .doc format from bytes requires file path on macOS"
                    result = subprocess.run(
                        ['textutil', '-convert', 'txt', '-stdout', file_path_or_bytes],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout
                    else:
                        return f"Error extracting .doc file: {result.stderr}"
                elif system == "Windows":
                    # Try using python-docx2txt or win32com
                    if is_bytes:
                        return "Error: .doc format from bytes requires file path on Windows"
                    try:
                        import docx2txt
                        return docx2txt.process(file_path_or_bytes)
                    except ImportError:
                        return "Error: .doc format requires docx2txt library on Windows"
                else:  # Linux
                    # Try antiword
                    if is_bytes:
                        return "Error: .doc format from bytes requires file path on Linux"
                    result = subprocess.run(
                        ['antiword', file_path_or_bytes],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    if result.returncode == 0:
                        return result.stdout
                    else:
                        return f"Error extracting .doc file: antiword not available or failed"
        except Exception as e:
            file_path_str = file_path_or_bytes if isinstance(file_path_or_bytes, str) else "<bytes>"
            logger.error(f"Error extracting Word document {file_path_str}: {e}", exc_info=True)
            return f"Error extracting Word document: {str(e)}"
    
    def _extract_excel(self, file_path_or_bytes) -> str:
        """Extract text from Excel file (.xls or .xlsx). Supports file path or BytesIO."""
        try:
            import pandas as pd
            import openpyxl  # For .xlsx support
            import io
            
            is_bytes = isinstance(file_path_or_bytes, (io.BytesIO, bytes))
            
            # Read all sheets
            if is_bytes:
                if isinstance(file_path_or_bytes, bytes):
                    file_path_or_bytes = io.BytesIO(file_path_or_bytes)
                excel_file = pd.ExcelFile(file_path_or_bytes)
            else:
                excel_file = pd.ExcelFile(file_path_or_bytes)
            
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"=== Sheet: {sheet_name} ===\n")
                # Convert DataFrame to text representation
                text_parts.append(df.to_string(index=False))
                text_parts.append("\n")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting Excel file: {e}", exc_info=True)
            return f"Error extracting Excel file: {str(e)}"
    
    def _extract_zip(self, file_path: str, extract_files: bool = False) -> str:
        """Extract content from ZIP archive."""
        try:
            text_parts = []
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                file_list = zip_ref.namelist()
                text_parts.append(f"ZIP Archive Contents ({len(file_list)} files):\n")
                text_parts.append("\n".join(f"  - {f}" for f in file_list))
                
                if extract_files:
                    text_parts.append("\n\n=== Extracted File Contents ===\n")
                    for file_name in file_list:
                        if file_name.endswith('/'):
                            continue  # Skip directories
                        
                        try:
                            # Read file content
                            content = zip_ref.read(file_name)
                            
                            # Try to extract text based on file extension
                            file_ext = os.path.splitext(file_name)[1].lower()
                            text_parts.append(f"\n--- {file_name} ---\n")
                            
                            if file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
                                # Text files
                                try:
                                    text_parts.append(content.decode('utf-8'))
                                except UnicodeDecodeError:
                                    text_parts.append(content.decode('latin-1'))
                            elif file_ext == '.pdf':
                                # Extract PDF text
                                import io
                                text_parts.append(self._extract_pdf(io.BytesIO(content)))
                            elif file_ext in ['.docx']:
                                # Extract Word text
                                import io
                                text_parts.append(self._extract_word(io.BytesIO(content)))
                            elif file_ext in ['.xlsx', '.xls']:
                                # Extract Excel text
                                import io
                                text_parts.append(self._extract_excel(io.BytesIO(content)))
                            else:
                                text_parts.append(f"[Binary file - {len(content)} bytes]")
                        except Exception as e:
                            text_parts.append(f"Error extracting {file_name}: {str(e)}")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting ZIP {file_path}: {e}", exc_info=True)
            return f"Error extracting ZIP archive: {str(e)}"
    
    def _extract_rar(self, file_path: str, extract_files: bool = False) -> str:
        """Extract content from RAR archive."""
        try:
            import rarfile
            
            text_parts = []
            with rarfile.RarFile(file_path, 'r') as rar_ref:
                file_list = rar_ref.namelist()
                text_parts.append(f"RAR Archive Contents ({len(file_list)} files):\n")
                text_parts.append("\n".join(f"  - {f}" for f in file_list))
                
                if extract_files:
                    text_parts.append("\n\n=== Extracted File Contents ===\n")
                    for file_name in file_list:
                        if file_name.endswith('/'):
                            continue  # Skip directories
                        
                        try:
                            # Read file content
                            content = rar_ref.read(file_name)
                            
                            # Try to extract text based on file extension
                            file_ext = os.path.splitext(file_name)[1].lower()
                            text_parts.append(f"\n--- {file_name} ---\n")
                            
                            if file_ext in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml']:
                                # Text files
                                try:
                                    text_parts.append(content.decode('utf-8'))
                                except UnicodeDecodeError:
                                    text_parts.append(content.decode('latin-1'))
                            elif file_ext == '.pdf':
                                # Extract PDF text
                                import io
                                text_parts.append(self._extract_pdf(io.BytesIO(content)))
                            elif file_ext in ['.docx']:
                                # Extract Word text
                                import io
                                text_parts.append(self._extract_word(io.BytesIO(content)))
                            elif file_ext in ['.xlsx', '.xls']:
                                # Extract Excel text
                                import io
                                text_parts.append(self._extract_excel(io.BytesIO(content)))
                            else:
                                text_parts.append(f"[Binary file - {len(content)} bytes]")
                        except Exception as e:
                            text_parts.append(f"Error extracting {file_name}: {str(e)}")
            
            return "\n".join(text_parts)
        except ImportError:
            return "Error: rarfile library not installed. Install with: pip install rarfile"
        except Exception as e:
            logger.error(f"Error extracting RAR {file_path}: {e}", exc_info=True)
            return f"Error extracting RAR archive: {str(e)}"
    
    def _run(self, file_path: str, extract_archives: bool = False, **kwargs) -> str:
        """Extract text content from the specified document file."""
        if not os.path.exists(file_path):
            return f"Error: File not found: {file_path}"
        
        if not os.path.isfile(file_path):
            return f"Error: Path is not a file: {file_path}"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._extract_word(file_path)
            elif file_ext in ['.xls', '.xlsx']:
                return self._extract_excel(file_path)
            elif file_ext == '.zip':
                return self._extract_zip(file_path, extract_archives)
            elif file_ext in ['.rar', '.rar5']:
                return self._extract_rar(file_path, extract_archives)
            elif self._is_text_file(file_path):
                # Handle text files (including unknown extensions that appear to be text)
                return self._extract_text_file(file_path)
            else:
                # Last resort: try to read as text anyway
                try:
                    return self._extract_text_file(file_path)
                except Exception:
                    return f"Error: Could not read file {file_path}. It may be a binary file."
        except Exception as e:
            logger.error(f"Error in document extraction: {e}", exc_info=True)
            return f"Error extracting document: {str(e)}"

