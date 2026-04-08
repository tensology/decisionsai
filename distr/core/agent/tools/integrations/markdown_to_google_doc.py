"""
Markdown to Google Doc Workflow Tool

This tool provides a seamless workflow to convert markdown content from clipboard
into a nicely formatted Google Doc, with automatic fallback to DOCX conversion
if direct creation fails. Prefers pandoc for MD→DOCX so pipe tables (GFM) are preserved.
"""

import logging
import re
import os
import subprocess
import platform
import shutil
import tempfile
from typing import Optional
from pathlib import Path
from datetime import datetime
from langchain.tools import BaseTool
from pydantic import BaseModel, Field
from distr.core.agent.services.integrations.google_workspace import GoogleWorkspaceConnector
from distr.core.agent.tools.base import LazyToolMixin

logger = logging.getLogger(__name__)

PANDOC_INSTALL_HINT = (
    "pandoc is not installed. For full table support install pandoc "
    "(macOS: brew install pandoc, Windows: choco install pandoc, Linux: apt install pandoc)."
)


def get_clipboard_content() -> Optional[str]:
    """Get content from clipboard using platform-specific methods."""
    try:
        system = platform.system()
        
        if system == "Darwin":  # macOS
            result = subprocess.run(
                ['pbpaste'],
                capture_output=True,
                text=True,
                timeout=1
            )
            return result.stdout if result.returncode == 0 else None
        elif system == "Windows":
            result = subprocess.run(
                ['powershell', '-command', 'Get-Clipboard'],
                capture_output=True,
                text=True,
                timeout=1
            )
            return result.stdout.strip() if result.returncode == 0 else None
        else:  # Linux
            try:
                result = subprocess.run(
                    ['xclip', '-selection', 'clipboard', '-o'],
                    capture_output=True,
                    text=True,
                    timeout=1
                )
                if result.returncode == 0:
                    return result.stdout
            except Exception:
                pass
            try:
                result = subprocess.run(
                    ['xsel', '--clipboard', '--output'],
                    capture_output=True,
                    text=True,
                    timeout=1
                )
                return result.stdout if result.returncode == 0 else None
            except Exception:
                pass
            return None
    except Exception as e:
        logger.error(f"Error getting clipboard content: {e}", exc_info=True)
        return None


def derive_title_from_markdown(markdown: str) -> str:
    """Derive title from first heading in markdown, or use default."""
    lines = markdown.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('#'):
            # Extract heading text
            title = stripped.lstrip('#').strip()
            if title:
                return title
    
    # Fallback: use first non-empty line (truncated) or default
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith('#'):
            # Take first 50 chars
            title = stripped[:50].strip()
            if title:
                return title
    
    # Ultimate fallback
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    return f"Document {timestamp}"


def _pandoc_available() -> bool:
    """Return True if pandoc is on PATH."""
    return shutil.which("pandoc") is not None


def convert_markdown_file_to_docx(
    input_md_path: str,
    output_path: Optional[str] = None,
) -> str:
    """
    Convert a Markdown file to DOCX using pandoc (GFM for table support).
    Output path defaults to same directory as input with .docx extension.
    Returns the path to the created DOCX file, or raises on failure.
    """
    path = Path(input_md_path).resolve()
    if not path.exists():
        raise FileNotFoundError(f"Markdown file not found: {input_md_path}")
    if output_path is None:
        output_path = str(path.with_suffix(".docx"))
    out = Path(output_path).resolve()
    pandoc_path = shutil.which("pandoc")
    if not pandoc_path:
        raise RuntimeError(PANDOC_INSTALL_HINT)
    result = subprocess.run(
        [pandoc_path, str(path), "-f", "gfm", "-t", "docx", "-o", str(out)],
        capture_output=True,
        text=True,
        timeout=60,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"pandoc failed: {result.stderr or result.stdout or 'unknown error'}"
        )
    logger.info("Converted markdown file to DOCX (pandoc GFM): %s -> %s", path, out)
    return str(out)


def _convert_markdown_to_docx_pandoc(markdown_content: str, output_path: str) -> bool:
    """
    Convert markdown to DOCX using pandoc with GFM input so pipe tables are preserved.
    Returns True on success, False on failure (e.g. pandoc not found or non-zero exit).
    """
    pandoc_path = shutil.which("pandoc")
    if not pandoc_path:
        logger.warning(PANDOC_INSTALL_HINT)
        return False
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            suffix=".md",
            delete=False,
            encoding="utf-8",
        ) as f:
            f.write(markdown_content)
            tmp_md = f.name
        try:
            result = subprocess.run(
                [pandoc_path, tmp_md, "-f", "gfm", "-t", "docx", "-o", output_path],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode != 0:
                logger.warning(
                    "pandoc MD→DOCX failed (exit %s): %s",
                    result.returncode,
                    result.stderr or result.stdout,
                )
                return False
            logger.info("Converted markdown to DOCX with pandoc (GFM, table support): %s", output_path)
            return True
        finally:
            try:
                os.unlink(tmp_md)
            except OSError:
                pass
    except Exception as e:
        logger.warning("pandoc conversion failed: %s", e)
        return False


def convert_markdown_to_docx(markdown_content: str, output_path: str) -> bool:
    """
    Convert markdown to DOCX. Prefers pandoc (GFM input) for full table support;
    falls back to python-docx if pandoc is unavailable or fails.
    """
    if _convert_markdown_to_docx_pandoc(markdown_content, output_path):
        return True
    # Fallback: python-docx (no pipe-table support)
    if not _pandoc_available():
        logger.info("Using python-docx fallback (no pandoc). %s", PANDOC_INSTALL_HINT)
    try:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Install with: pip install python-docx"
            ) from e

        doc = Document()
        lines = markdown_content.split('\n')
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Headings
            if stripped.startswith('#'):
                level = len(stripped) - len(stripped.lstrip('#'))
                text = stripped.lstrip('#').strip()
                if text:
                    heading = doc.add_heading(text, level=min(level, 6))
                    heading.style.font.size = Pt(24 - (level * 2))
            
            # Horizontal rules
            elif stripped in ['---', '***', '___']:
                doc.add_paragraph('─' * 50)
            
            # Unordered lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:].strip()
                # Remove markdown formatting
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'\*([^*]+)\*', r'\1', text)
                doc.add_paragraph(text, style='List Bullet')
            
            # Ordered lists
            elif re.match(r'^\d+\.\s+', stripped):
                text = re.sub(r'^\d+\.\s+', '', stripped)
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'\*([^*]+)\*', r'\1', text)
                doc.add_paragraph(text, style='List Number')
            
            # Bold text (standalone line)
            elif '**' in stripped or '__' in stripped:
                text = stripped
                # Replace markdown bold
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'__([^_]+)__', r'\1', text)
                p = doc.add_paragraph(text)
                # Apply bold to the paragraph
                for run in p.runs:
                    run.bold = True
            
            # Regular paragraph
            elif stripped:
                # Remove markdown formatting
                text = stripped
                text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
                text = re.sub(r'\*([^*]+)\*', r'\1', text)
                text = re.sub(r'`([^`]+)`', r'\1', text)
                text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
                doc.add_paragraph(text)
            
            # Empty line
            else:
                doc.add_paragraph()
            
            i += 1
        
        doc.save(output_path)
        logger.info(f"Successfully converted markdown to DOCX: {output_path}")
        return True
        
    except RuntimeError:
        raise
    except Exception as e:
        logger.error(f"Failed to convert markdown to DOCX: {e}", exc_info=True)
        return False


class MarkdownToGoogleDocInput(BaseModel):
    """Input schema for markdown to Google Doc tool."""
    open_in_brave: bool = Field(default=True, description="Whether to open the document in Brave browser after creation")
    use_fallback: bool = Field(default=False, description="Force use of fallback path (DOCX conversion) instead of direct creation")


class MarkdownToGoogleDocTool(LazyToolMixin, BaseTool):
    """Tool for converting markdown from clipboard to Google Doc with automatic fallback."""
    
    name: str = "markdown_to_google_doc"
    description: str = (
        "Convert markdown content from clipboard to a nicely formatted Google Doc.\n"
        "This tool automatically:\n"
        "- Gets markdown content from clipboard\n"
        "- Derives title from first heading or uses default\n"
        "- Creates Google Doc with proper formatting (headings, bold, lists, horizontal rules)\n"
        "- Falls back to DOCX conversion if direct creation fails\n"
        "- Opens the document in Brave browser\n"
        "\n"
        "Path A (Primary): Direct Google Doc creation from markdown with formatting\n"
        "Path B (Fallback): Convert markdown to DOCX locally, upload to Drive, convert to Google Doc\n"
        "\n"
        "Use this when user says:\n"
        "- 'Take what's in my clipboard, it's markdown, convert it to a Google Doc and open it in Brave'\n"
        "- 'Create a Google Doc from my clipboard markdown'\n"
        "- 'Convert my clipboard to a Google document'\n"
    )
    args_schema: type[BaseModel] = MarkdownToGoogleDocInput
    
    def __init__(self):
        super().__init__()

    def _lazy_init(self):
        object.__setattr__(self, 'connector', GoogleWorkspaceConnector())
    
    def _run(self, open_in_brave: bool = True, use_fallback: bool = False, **kwargs) -> str:
        """Execute markdown to Google Doc conversion workflow"""
        self._ensure_initialized()
        try:
            # Step 1: Get clipboard content
            logger.info("Getting clipboard content...")
            markdown_content = get_clipboard_content()
            
            if not markdown_content or not markdown_content.strip():
                return "Error: Clipboard is empty. Please copy markdown content to clipboard first."
            
            logger.info(f"Got clipboard content ({len(markdown_content)} chars)")
            
            # Step 2: Derive title
            title = derive_title_from_markdown(markdown_content)
            logger.info(f"Derived title: {title}")
            
            # Check if Google is connected
            if not self.connector.is_connected():
                return "Error: Google is not connected. Please connect your Google account in Settings > Advanced."
            
            # Step 3: Try Path A (Direct creation) unless fallback is forced
            if not use_fallback:
                logger.info("Attempting Path A: Direct Google Doc creation from markdown...")
                doc_id = self.connector.create_doc_from_markdown(
                    title=title,
                    markdown_content=markdown_content,
                    folder_id="root",
                    preserve_formatting=True
                )
                
                if doc_id:
                    doc_url = self.connector.get_document_url(doc_id)
                    logger.info(f"Successfully created Google Doc: {doc_id}, URL: {doc_url}")
                    
                    # Step 4: Open in Brave if requested
                    if open_in_brave:
                        self.connector.open_url_in_brave(doc_url)
                    
                    return f"Successfully created Google Doc '{title}' (ID: {doc_id}). Document URL: {doc_url}"
                else:
                    logger.warning("Path A failed, falling back to Path B...")
            
            # Step 5: Path B (Fallback) - DOCX conversion
            logger.info("Using Path B: DOCX conversion fallback...")
            
            # Save markdown to file
            downloads_dir = Path.home() / "Downloads"
            timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
            markdown_file = downloads_dir / f"markdown-{timestamp}.md"
            
            try:
                with open(markdown_file, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                logger.info(f"Saved markdown to: {markdown_file}")
            except Exception as e:
                logger.error(f"Failed to save markdown file: {e}", exc_info=True)
                return f"Error: Failed to save markdown file: {e}"
            
            # Convert to DOCX
            docx_file = downloads_dir / f"{title}-{timestamp}.docx"
            # Sanitize filename
            safe_title = re.sub(r'[^\w\s-]', '', title)[:50]
            docx_file = downloads_dir / f"{safe_title}-{timestamp}.docx"
            
            if not convert_markdown_to_docx(markdown_content, str(docx_file)):
                return "Error: Failed to convert markdown to DOCX. Please ensure python-docx is installed."
            
            logger.info(f"Converted to DOCX: {docx_file}")
            
            # Upload DOCX to Drive with conversion
            logger.info("Uploading DOCX to Google Drive with conversion...")
            file_id = self.connector.upload_to_drive(
                file_path=str(docx_file),
                folder_id="root",
                name=safe_title,
                convert_to_google_doc=True
            )
            
            if not file_id:
                # Try uploading without conversion, then convert separately
                logger.info("Upload without conversion failed, trying separate upload and convert...")
                file_id = self.connector.upload_to_drive(
                    file_path=str(docx_file),
                    folder_id="root",
                    name=safe_title,
                    convert_to_google_doc=False
                )
                
                if file_id:
                    # Convert to Google Doc
                    google_doc_id = self.connector.convert_docx_to_google_doc(file_id)
                    if google_doc_id:
                        file_id = google_doc_id
                    else:
                        return f"Error: Uploaded DOCX (ID: {file_id}) but failed to convert to Google Doc. You can access it at: https://drive.google.com/file/d/{file_id}/view"
            
            if not file_id:
                return "Error: Failed to upload DOCX to Google Drive."
            
            doc_url = self.connector.get_document_url(file_id)
            logger.info(f"Successfully created Google Doc via fallback: {file_id}, URL: {doc_url}")
            
            # Clean up local files
            try:
                if markdown_file.exists():
                    markdown_file.unlink()
                if docx_file.exists():
                    docx_file.unlink()
            except Exception as e:
                logger.warning(f"Failed to clean up local files: {e}")
            
            # Open in Brave if requested
            if open_in_brave:
                self.connector.open_url_in_brave(doc_url)
            
            return f"Successfully created Google Doc '{title}' via fallback method (ID: {file_id}). Document URL: {doc_url}"
            
        except Exception as e:
            logger.error(f"Error in markdown to Google Doc workflow: {e}", exc_info=True)
            return f"Error: {str(e)}"
    
    async def _arun(self, open_in_brave: bool = True, use_fallback: bool = False, **kwargs) -> str:
        """Async run method"""
        self._ensure_initialized()
        return self._run(open_in_brave=open_in_brave, use_fallback=use_fallback, **kwargs)

